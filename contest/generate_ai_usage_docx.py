from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = Path("/Users/talexdreamsoul/Downloads/4-AI工具使用说明.docx")
OUTPUT_DOCX = ROOT / "AI工具使用说明-WinLoop.docx"
README = ROOT / "README.md"

PROJECT_ID = "待填写"
PROJECT_NAME = "WinLoop AI（赛帮帮）- 面向竞赛团队的智能协作与答辩平台"
USAGE_PERIOD = "2026年3月1日-5月5日"


HEADERS = [
    "序号",
    "AI工具的名称、版本、访问方式（网页、API或客户端），使用时间",
    "使用AI工具的环节与目的（立项构思、文献综述、语言润色、内容生成、图表优化、代码编程、数据分析等）",
    "关键提示词",
    "AI回复的关键内容（在此简要说明，并在附录中给出佐证）",
    "AI回复的\n人工修改说明",
    "采纳比例\n与说明",
]


ROWS = [
    [
        "1",
        f"通义千问 Qwen-Plus / Qwen-Max（大语言模型）\nAPI访问：百炼 DashScope OpenAI兼容接口\n使用时间：{USAGE_PERIOD}",
        "内容生成、资料问答、项目分析：用于工作台右侧助手、项目聊天和竞赛资料问答。模型结合项目资源、协作文档、比赛信息和知识索引上下文，辅助团队理解资料、梳理项目状态、生成下一步建议。",
        "System Prompt: \"你是WinLoop竞赛团队工作台助手。请基于当前项目资料、比赛信息和可见引用回答，不确定时说明资料缺口，不编造结论。\"\nUser示例: \"根据当前资料，帮我整理项目亮点、风险和下一步任务。\"",
        "1. 输出项目亮点、资料缺口、风险项和下一步建议。\n2. 能结合竞赛绑定、项目设置、资源列表和协作文档摘要回答。\n3. 在知识索引未完成时提示结果可能不完整，避免把不确定内容写成确定结论。",
        "人工保留了项目权限、资源可见性和引用状态的校验；对模型输出进行了事实核对、术语统一和中文表达润色；涉及写操作的建议仍通过草案或人工确认，不直接落盘。",
        "采纳比例约80%。主要采纳结构化梳理、说明文本和问答结果；项目事实、权限边界、资料引用和最终表达由人工复核后确定。",
    ],
    [
        "2",
        f"通义千问 Qwen-VL-Max / Qwen-VL-Plus（多模态视觉理解模型）\nAPI访问：百炼 DashScope 多模态接口\n使用时间：{USAGE_PERIOD}",
        "内容生成、数据分析、图像理解：用于图片资料、OCR文本、流程图截图、设计稿截图和会议视觉材料的理解，辅助生成项目知识片段和可检索摘要。",
        "Prompt示例: \"请阅读这张竞赛资料截图，提取标题、关键结论、时间节点和可能影响项目答辩的证据点。输出为结构化JSON，保留不确定项。\"\nPrompt示例: \"请概括这张流程画布截图中的节点关系，用于项目知识索引。\"",
        "1. 生成图片摘要、OCR文本、流程图节点说明和视觉资料标签。\n2. 将视觉内容投影到统一文本语义空间，进入项目知识索引。\n3. 输出可供工作台AI引用的资料片段，支持后续问答和答辩准备。",
        "人工补充了资料来源、项目归属和引用位置；删除低置信度识别结果；对OCR误字、图中缩写和流程节点命名进行校正，避免视觉误判进入正式材料。",
        "采纳比例约75%。主要采纳视觉摘要和结构化提取结果；截图来源、项目术语、引用范围和最终入库内容由人工确认。",
    ],
    [
        "3",
        f"通义千问 Qwen Realtime / Qwen ASR / Qwen TTS（实时音视频、语音识别、语音合成能力）\nAPI访问：百炼 DashScope WebSocket / 语音接口\n使用时间：{USAGE_PERIOD}",
        "内容生成、语音转写、答辩模拟：用于答辩工作台中的实时答辩sidecar、语音转文字、答辩问题理解、实时建议和语音播报。系统保留LiveKit会议链路，Qwen负责实时推理和语音能力。",
        "Realtime Prompt示例: \"你是竞赛答辩评委模拟助手。请根据当前项目资料、发言转写和评委persona提出追问，并给出简短改进建议。\"\nASR/TTS配置示例: \"中文普通话实时转写，保留中间结果；播报时语气清晰、适合答辩场景。\"",
        "1. 支持答辩会话中的实时转写和评委追问。\n2. 将turn、摘要、问题和动作项回流到答辩工作台。\n3. 对AI未配置、连接失败和权限问题给出明确状态，不静默降级。",
        "人工调整了答辩persona、问题强度、转写分段和摘要粒度；对实时建议进行筛选，只保留可用于答辩训练和会后复盘的内容。",
        "采纳比例约70%。实时转写和初稿建议可直接辅助训练，但最终答辩话术、评分判断和会后总结需要人工再加工。",
    ],
    [
        "4",
        f"DeepSeek V3 / DeepSeek R1（大语言模型、推理模型）\n访问方式：网页端与OpenAI兼容API\n使用时间：{USAGE_PERIOD}",
        "立项构思、方案推演、文档润色：用于梳理竞赛团队痛点、产品定位、功能边界、评审问题和技术讲解顺序，辅助把WinLoop从功能列表整理为完整参赛叙事。",
        "Prompt示例: \"请基于竞赛团队从选赛、建项、资料沉淀、协作推进到答辩提交的流程，梳理一个5分钟技术路演结构。要求突出多模态知识索引、可溯源引用和答辩工作台，不要泛泛讲AI。\"",
        "1. 输出痛点、产品定位、核心链路、技术亮点和评审追问清单。\n2. 将复杂能力归纳为竞赛协作、知识索引、智能工作流、答辩辅助和后台治理五类。\n3. 给出更适合比赛材料的表达顺序。",
        "人工根据当前仓库真实实现删去了未落地或表述过度的内容；将空泛AI词汇改为多模态语义索引、混合召回、可信引用、实时答辩等可解释能力。",
        "采纳比例约65%。主要采纳结构和问题清单；项目事实、实现边界和正式文案均由人工结合仓库资料重写。",
    ],
    [
        "5",
        f"DeepSeek V3 + Qwen-Plus（文本生成与推理模型组合）\n访问方式：网页端、API与工作台模型路由\n使用时间：{USAGE_PERIOD}",
        "需求验收、测试点生成、异常分析：用于对项目知识索引、AI provider治理、飞书同步、答辩工作台等模块做验收清单、边界条件和异常说明整理。仅少量辅助接口说明和测试用例梳理，不作为主要代码生成工具。",
        "Prompt示例: \"请把以下功能说明整理成验收清单：项目资料上传后进入知识索引，ready优先，stale仅fallback，AI回答要展示citation和warning。输出正常路径、异常路径和验收证据。\"",
        "1. 生成接口验收点、失败路径、状态文案和回归测试建议。\n2. 帮助识别未配置provider、索引未完成、资料不可见、引用缺失等风险。\n3. 为文档中的人工修改说明提供依据。",
        "人工将模型输出压缩为实际可验证的场景；删除与当前仓库不一致的接口命名；测试命令和验收结论以本地项目实际结果为准。",
        "采纳比例约60%。主要采纳测试思路和异常清单；具体实现、命名、断言和验收结论由人工确认。",
    ],
    [
        "6",
        f"即梦AI 文生图（图像生成工具）\n访问方式：网页端/客户端\n使用时间：{USAGE_PERIOD}",
        "内容生成、图表优化、比赛包装：用于生成WinLoop产品封面、路演主视觉、竞赛团队工作台氛围图和申报材料配图草案，增强比赛材料的视觉表达。",
        "Prompt示例: \"面向大学生竞赛团队的智能协作平台，画面包含项目工作台、知识网络、答辩会议和资料卡片，现代中文SaaS产品视觉，干净、专业、科技感，16:9，适合路演封面，不要出现具体品牌Logo。\"",
        "1. 生成多版产品主视觉和封面草案。\n2. 提供知识网络、工作台协作、答辩辅助等主题的视觉方向。\n3. 为PPT和申报材料选择更统一的视觉基调。",
        "人工筛选构图，统一品牌色和标题层级；删除不符合产品实际的装饰元素；对文字区域留白、人物姿态、界面细节进行二次修正。",
        "采纳比例约70%。主要采纳构图、氛围和视觉草案；最终版面、标题、配色和是否用于提交由人工决定。",
    ],
    [
        "7",
        f"即梦AI 图生图 / 风格化（图像生成与二次处理工具）\n访问方式：网页端/客户端\n使用时间：{USAGE_PERIOD}",
        "图表优化、UI展示、功能示意：用于对工作台截图、流程画布截图和答辩界面草图做风格统一、背景扩展、视觉包装和演示图优化。",
        "Prompt示例: \"基于这张工作台截图，保持原有布局和信息层级，优化为比赛路演展示图。风格简洁专业，不添加不存在的功能按钮，不改变中文文案。\"",
        "1. 生成更适合路演展示的工作台示意图。\n2. 对截图边缘、背景和光影进行包装，突出协作、资料和答辩状态。\n3. 帮助制作功能页面之间的视觉过渡图。",
        "人工核对所有界面元素，避免模型误改按钮、数字、状态和中文文案；真实产品截图仍作为验收依据，AI图只用于视觉包装或示意说明。",
        "采纳比例约55%。主要采纳视觉包装效果；涉及真实界面、数据和功能状态的内容必须人工复核。",
    ],
    [
        "8",
        f"即梦AI 素材扩展（图标、背景、插画与配图生成）\n访问方式：网页端/客户端\n使用时间：{USAGE_PERIOD}",
        "内容生成、素材制作：用于生成答辩材料中的章节背景、能力图标、知识索引插画、评委问答场景图和社交传播配图草案。",
        "Prompt示例: \"生成一组用于竞赛协作平台PPT的扁平化图标：知识索引、可信引用、团队协作、答辩模拟、飞书同步、AI治理。统一线性风格，浅色背景，图标本身不包含文字。\"",
        "1. 输出多组图标和插画方向。\n2. 帮助统一比赛材料中的章节视觉。\n3. 用于辅助说明抽象能力，例如知识图谱、引用链路和实时答辩。",
        "人工统一尺寸、留白、对齐和色彩；删除与项目无关或过度装饰的图形；最终图标按材料版式重新排版。",
        "采纳比例约65%。主要采纳图形方向和素材草案；最终视觉规范、语义匹配和排版由人工完成。",
    ],
    [
        "9",
        f"Qwen-Plus / DeepSeek V3（文本整理模型）\n访问方式：网页端与API\n使用时间：{USAGE_PERIOD}",
        "内容生成、资料整理、会议纪要：用于飞书多维表导入字段说明、项目资料摘要、会议纪要整理、答辩动作项归纳和后台运营说明润色。",
        "Prompt示例: \"请将这些竞赛记录整理为同步摘要，区分飞书原始字段、本地保留字段、发布后保留字段，并列出需要人工确认的缺失信息。\"",
        "1. 输出飞书同步摘要、资料缺口和人工确认项。\n2. 将会议转写整理为总结、问题、动作项和证据缺口。\n3. 辅助后台运营材料说明AI provider、资源worker和同步状态。",
        "人工核验飞书字段来源、同步边界和会议事实；对模型生成的总结进行删改，保证不把未确认事项写成已完成事项。",
        "采纳比例约70%。资料整理和纪要结构较高比例采纳；字段来源、状态判断和最终提交内容由人工确认。",
    ],
]


APPENDIX_NOTE = """填写说明：
本文档适用于所有涉及AI工具使用的参赛作品，表格可另加行。
参赛作品的作者，需根据实际的使用情况简明扼要地列出本作品所使用的全部AI工具的名称、版本、访问方式、使用时间、使用环节与目的、关键提示词、AI回复的关键内容、采纳和人工修改情况等。
AI回复的关键内容佐证材料，需作为本文档的附录2给出，包括但不限于：（1）关键操作截图（含时间戳，需清晰可辨）；（2）交互录屏视频（时长≤5分钟，需标注使用节点，文档为MP4格式，命名格式：AI_使用序号_作品编号.mp4）；（3）代码注释中标明AI辅助部分（如：// AI辅助生成：DeepSeek-R1-0528, 2026-05-05）
提交时，需将本文档的PDF格式文件，以及其他佐证材料（如交互录屏视频），一并上传到作品文件夹的“03设计与开发文档”子文件夹中。
本文档内容是正式参赛内容组成部分，需真实填写。如不属实，将导致奖项等级降低甚至终止本作品参加比赛。"""


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 8.0) -> None:
    cell.text = ""
    lines = text.split("\n")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)


def set_paragraph_text(paragraph, text: str, *, font_size: float = 10.5, bold: bool = False, align=None) -> None:
    paragraph.text = ""
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)


def remove_extra_rows(table, keep_rows: int) -> None:
    while len(table.rows) > keep_rows:
        row = table.rows[-1]
        row._element.getparent().remove(row._element)


def ensure_row_count(table, row_count: int) -> None:
    while len(table.rows) < row_count:
        table.add_row()


def set_column_widths(table) -> None:
    widths = [Cm(0.9), Cm(4.2), Cm(4.8), Cm(5.2), Cm(4.6), Cm(4.4), Cm(3.7)]
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = width


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_vertical_alignment(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "top")


def format_table(table) -> None:
    set_column_widths(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_vertical_alignment(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
            if row_index == 0:
                shade_cell(cell, "D9EAF7")


def trim_paragraphs_after_table(document: Document) -> None:
    for paragraph in list(document.paragraphs[3:]):
        paragraph._element.getparent().remove(paragraph._element)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"模板不存在：{SOURCE_DOCX}")

    document = Document(SOURCE_DOCX)

    section = document.sections[0]
    section.left_margin = Cm(0.9)
    section.right_margin = Cm(0.9)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    set_paragraph_text(document.paragraphs[0], "中国大学生计算机设计大赛", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(document.paragraphs[1], "AI工具使用说明(2026年版)", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_text(
        document.paragraphs[2],
        f"作品编号：  {PROJECT_ID}       作品名称：    {PROJECT_NAME}",
        font_size=10.5,
        bold=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    table = document.tables[0]
    ensure_row_count(table, len(ROWS) + 1)
    remove_extra_rows(table, len(ROWS) + 1)

    for index, header in enumerate(HEADERS):
        set_cell_text(table.rows[0].cells[index], header, bold=True, font_size=8.0)

    for row_index, row_data in enumerate(ROWS, start=1):
        for col_index, text in enumerate(row_data):
            set_cell_text(table.rows[row_index].cells[col_index], text, font_size=7.4 if col_index else 8.0)

    format_table(table)
    trim_paragraphs_after_table(document)

    document.add_paragraph()
    set_paragraph_text(document.add_paragraph(), APPENDIX_NOTE, font_size=9.0)

    add_page_break(document)
    set_paragraph_text(document.add_paragraph(), "附录1：作品文件夹示例", font_size=11, bold=True)
    folder_lines = [
        f"{PROJECT_ID}-参赛总文件夹",
        f"├── {PROJECT_ID}-01作品与答辩材料",
        f"├── {PROJECT_ID}-02素材与源码",
        f"├── {PROJECT_ID}-03设计与开发文档",
        f"└── {PROJECT_ID}-04作品演示视频",
    ]
    for line in folder_lines:
        set_paragraph_text(document.add_paragraph(), line, font_size=9.5)

    document.add_paragraph()
    set_paragraph_text(document.add_paragraph(), "附录2：AI回复关键内容佐证材料清单（待补）", font_size=11, bold=True)
    for row in ROWS:
        paragraph = document.add_paragraph()
        set_paragraph_text(
            paragraph,
            f"序号{row[0]}的佐证材料：待补关键操作截图或录屏，建议命名为 AI_{row[0]}_{PROJECT_ID}.mp4；截图文件可命名为 AI_{row[0]}_{PROJECT_ID}_截图.png。",
            font_size=9.5,
        )

    document.save(OUTPUT_DOCX)


def write_readme() -> None:
    README.write_text(
        "\n".join(
            [
                "# WinLoop AI工具使用说明",
                "",
                f"- 生成日期：{date.today().isoformat()}",
                f"- 模板来源：`{SOURCE_DOCX}`",
                f"- 输出文件：`{OUTPUT_DOCX.name}`",
                f"- 作品编号：`{PROJECT_ID}`，提交前需要替换为正式编号。",
                f"- 作品名称：`{PROJECT_NAME}`",
                "",
                "## 后续需要补充",
                "",
                "1. 将 Word 和导出的 PDF 放入参赛作品文件夹的 `03设计与开发文档`。",
                "2. 按附录2补齐每个 AI 工具使用序号对应的截图或录屏佐证。",
                "3. 若正式作品名称或编号变化，更新本目录脚本中的 `PROJECT_ID` / `PROJECT_NAME` 后重新生成。",
                "",
                "## 生成命令",
                "",
                "```bash",
                '"/Users/talexdreamsoul/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3" "contest/generate_ai_usage_docx.py"',
                "```",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_document()
    write_readme()
    print(OUTPUT_DOCX)
