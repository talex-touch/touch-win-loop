from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "12-application-material.md"
OUTPUT_DIR = ROOT / "exports"
OUTPUT_DOCX = OUTPUT_DIR / "WinLoop-正式申报材料.docx"

IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>[^)]+)\)")


def set_run_font(run, font_name: str, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, *, font_name: str = "Microsoft YaHei", size_pt: float = 11.5, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt=size_pt, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(header, "WinLoop 正式申报材料", size_pt=9.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(footer, "第 ", size_pt=9.5)
    add_page_number(footer)
    set_paragraph_text(footer, " 页", size_pt=9.5)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WinLoop 正式申报材料")
    set_run_font(run, "Microsoft YaHei", size_pt=22, bold=True)
    p.space_after = Pt(14)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("竞赛团队的智能作战工作台")
    set_run_font(run2, "Microsoft YaHei", size_pt=15)
    p2.space_after = Pt(20)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(
        p3,
        "内容基于当前仓库实现整理，适用于正式申报、评审归档与比赛提交。",
        size_pt=11.5,
    )

    document.add_page_break()


def flush_paragraph(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines).strip()
    if not text:
        return
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_text(p, text)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    size_map = {1: 18, 2: 15, 3: 13}
    set_run_font(run, "Microsoft YaHei", size_pt=size_map.get(level, 13), bold=True)


def add_quote(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, "Microsoft YaHei", size_pt=10.5)
    run.italic = True


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_text(p, text)


def add_image(document: Document, alt: str, rel_path: str) -> None:
    image_path = (SOURCE_MD.parent / rel_path).resolve()
    if not image_path.exists():
        p = document.add_paragraph()
        set_paragraph_text(p, f"[图片缺失] {alt} - {rel_path}", size_pt=10.5)
        return

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(16.2))

    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(cap, alt or image_path.stem, size_pt=9.5)


def parse_markdown_to_docx(document: Document, text: str) -> None:
    buffered: list[str] = []
    skipped_cover_title = False

    def flush() -> None:
        nonlocal buffered
        flush_paragraph(document, buffered)
        buffered = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush()
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush()
            add_image(document, image_match.group("alt"), image_match.group("path"))
            continue

        if stripped.startswith("#"):
            flush()
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped[level:].strip()
            if level == 1 and not skipped_cover_title and heading_text == "WinLoop 正式申报材料":
                skipped_cover_title = True
                continue
            add_heading(document, heading_text, level)
            continue

        if stripped.startswith("> "):
            flush()
            add_quote(document, stripped[2:].strip())
            continue

        if stripped.startswith("- "):
            flush()
            add_bullet(document, stripped[2:].strip())
            continue

        buffered.append(stripped)

    flush()


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    source_text = SOURCE_MD.read_text(encoding="utf-8")

    document = Document()
    configure_document(document)
    add_cover(document)
    parse_markdown_to_docx(document, source_text)

    document.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
